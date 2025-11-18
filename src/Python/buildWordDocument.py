from pathlib import Path
import argparse
import json

from docxtpl import DocxTemplate
from docx import Document
from docx.enum.text import WD_COLOR_INDEX


# ---------- Step 1: Build context for docxtpl ----------

def build_context_from_json(data):
    """
    Convert JSON keys like "{ESPBoldName}" into Jinja variable names like "ESPBoldName".
    """
    context = {}
    for key, value in data.items():
        if key.startswith("{") and key.endswith("}"):
            clean_key = key[1:-1]  # strip surrounding braces
        else:
            clean_key = key
        context[clean_key] = value
    return context


# ---------- Step 2: Post-process rendered docx to highlight/bold replacements ----------

def highlight_replacements(doc_path, context, bold_keys=None):
    """
    Open the rendered DOCX and highlight all replacement values in yellow.
    Also bold those whose key is in bold_keys.
    """
    if bold_keys is None:
        bold_keys = set()

    doc = Document(doc_path)

    # We'll search for the replacement values (context values) in runs.
    # This assumes each placeholder renders as its own run (which is typical
    # when using {{ var }} in a template).
    def process_runs(runs):
        for run in runs:
            text = run.text
            if not text:
                continue

            for key, value in context.items():
                if value is None:
                    continue
                value_str = str(value)
                if value_str and value_str == text:
                    # Highlight this run (it's exactly the replacement value)
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

                    # Bold if this key should be bold
                    if key in bold_keys:
                        run.font.bold = True
                    break  # done with this run

    # Process paragraphs
    for para in doc.paragraphs:
        process_runs(para.runs)

    # Process tables as well (if any placeholders are inside tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_runs(para.runs)

    doc.save(doc_path)


# ---------- Step 3: Wire everything together ----------

def render_docx_template(template_path, output_path, data, bold_keys=None):
    # Build the Jinja context from JSON keys
    context = build_context_from_json(data)

    # Render with docxtpl
    doc = DocxTemplate(str(template_path))
    doc.render(context)
    doc.save(str(output_path))

    # After rendering, highlight all changes and bold selected ones
    highlight_replacements(output_path, context, bold_keys=bold_keys)

    print("\033[32mSuccessfully saved document\033[0m")


# ---------- CLI handling ----------

parser = argparse.ArgumentParser(description="Create a warrant with existing template.")
parser.add_argument("template", help="Path to the template .docx file.")
parser.add_argument("saveDestination", help="Path of save directory.")
parser.add_argument(
    "outputName",
    nargs="?",
    default="Warrant_Output.docx",
    help="Name of the output .docx file (default: Warrant_Output.docx).",
)
args = parser.parse_args()

TEMPLATE_PATH = Path(args.template)
if not TEMPLATE_PATH.is_file():
    raise SystemExit(f"\033[41mNot a valid template file:\033[0m {TEMPLATE_PATH.resolve()}")

SAVE_PATH = Path(args.saveDestination)
if not SAVE_PATH.is_dir():
    raise SystemExit(f"\033[41mNot a valid save path:\033[0m {SAVE_PATH.resolve()}")

OUTPUT_PATH = SAVE_PATH / args.outputName

with open('resources/results.json', 'r') as file:
    data = json.load(file)

try:
    # These must match the variable names you use in the template: {{ ESPBoldName }}
    bold_keys = {"ESPBoldName"}
    render_docx_template(TEMPLATE_PATH, OUTPUT_PATH, data, bold_keys=bold_keys)
except Exception as e:
    print(f"\033[31mSomething went wrong: {e}\033[0m")
